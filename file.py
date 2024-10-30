from message import print_alert_message, print_warning_message
from validate import is_image
from pypdf import PdfReader
from PIL import Image

from ai import get_new_file_name, get_new_file_name_for_image

import docx
import random
import os
import json

supported_extensions = [".txt", ".pdf", ".docx"]
supported_image_extensions = [ext for ext, format_ in Image.registered_extensions().items() if format_ in Image.OPEN]

class File:
    file_path = ""
    file_extension = ""
    file_exists = True
    def __init__(self,file_path):
        if os.path.isfile(file_path):
            self.file_extension = os.path.splitext(file_path)[1]
            self.file_path = file_path
        else:
            print_warning_message("The file does not exist.")
            self.file_exists = False

    def _get_random_page_number(self, stop_at):
        return random.randrange(0, stop_at)
    def _read_pdf(self):
        # creating a pdf reader object
        reader = PdfReader(self.file_path)
        text = reader.pages[0].extract_text()
        text += reader.pages[
            self._get_random_page_number(len(reader.pages))
        ].extract_text()
        text += reader.pages[
            self._get_random_page_number(len(reader.pages))
        ].extract_text()

        return text
    def _read_txt(self):
        with open(self.file_path, 'r') as file:
            data = file.read()
        return data

    def _get_docx(self):
        doc = docx.Document(self.file_path)
        fullText = doc.paragraphs[0].text
        fullText += doc.paragraphs[
            self._get_random_page_number(len(doc.paragraphs))
        ].text
        fullText += doc.paragraphs[
            self._get_random_page_number(len(doc.paragraphs))
        ].text
        fullText += doc.paragraphs[
            self._get_random_page_number(len(doc.paragraphs))
        ].text
        fullText += doc.paragraphs[
            self._get_random_page_number(len(doc.paragraphs))
        ].text
        return fullText

    def _read(self):
        text = ""
        #To prevent printing message "Unfortunately, we do not support this file extension."
        if self.file_exists:
            if self.file_extension == ".txt":
                text = self._read_txt()
            elif self.file_extension == ".pdf":
                text = self._read_pdf()
            elif self.file_extension == ".docx":
                text = self._get_docx()
            if is_image(self.file_path):
                text = "image"
            else:
                print_warning_message("Unfortunately, we do not support this file extension.")
        return text
    def rename(self):
        text = self._read()
        if text != "":
            #To convert a string containing a JSON object into a real JSON object in Python
            response = {}
            if is_image(self.file_path):
                response = {"filename":get_new_file_name_for_image(self.file_path)}
            else:
                response = json.loads(get_new_file_name(text))
            new_file_name = f"{response['filename']}.{self.file_extension}"
            #Get directory without file basename
            directory = os.path.dirname(self.file_path)
            new_file = os.path.join(directory, new_file_name)
            try:
                os.rename(self.file_path, new_file)
                print_warning_message("The file was renamed")
            except:
                print_alert_message("An exception occurred")